from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph
)

from reportlab.lib.styles import (
    getSampleStyleSheet
)

from docx import Document


def export_txt(filepath, messages):

    with open(
        filepath,
        "w",
        encoding="utf-8"
    ) as f:

        for role, content, _ in messages:

            f.write(
                f"{role}: {content}\n\n"
            )


def export_docx(filepath, messages):

    doc = Document()

    doc.add_heading(
        "Teja.AI Chat Export",
        0
    )

    for role, content, _ in messages:

        doc.add_paragraph(
            f"{role}: {content}"
        )

    doc.save(filepath)


def export_pdf(filepath, messages):

    pdf = SimpleDocTemplate(
        filepath
    )

    styles = getSampleStyleSheet()

    story = []

    for role, content, _ in messages:

        story.append(
            Paragraph(
                f"<b>{role}</b>: {content}",
                styles["BodyText"]
            )
        )

    pdf.build(story)